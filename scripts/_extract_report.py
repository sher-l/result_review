"""提取 25YHB676F 报告文本并准备审核"""
import sys
from pathlib import Path

proj_root = Path(r"D:\IKL\BaiduSyncdisk\报告审核\raw\待审核\25YHB676F-骨质疏松的遗传学研究：多组学方法（转录组单细胞以及孟德尔随机化）识别潜在治疗靶点--pzq-2025.2.02")

# 找 docx
docx_files = list(proj_root.glob("*.docx"))
if not docx_files:
    print("ERROR: No docx found")
    sys.exit(1)

docx_path = docx_files[0]
print(f"DOCX: {docx_path.name}")

import docx
doc = docx.Document(str(docx_path))

# 提取段落文本
lines = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        lines.append(text)

# 提取表格文本
for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append(" | ".join(cells))

report_text = "\n".join(lines)
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Text length: {len(report_text)} chars")

# 保存 report_text.txt
out_path = proj_root / "report_text.txt"
out_path.write_text(report_text, encoding="utf-8")
print(f"Saved: {out_path}")

# 也拷贝到 result_review_report
review_dir = Path(r"d:\IKL\BaiduSyncdisk\报告审核\result_review_report\25YHB676F")
review_dir.mkdir(parents=True, exist_ok=True)
(review_dir / "report_text.txt").write_text(report_text, encoding="utf-8")
print(f"Copied to: {review_dir / 'report_text.txt'}")
